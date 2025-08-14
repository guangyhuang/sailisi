from langchain_chroma import Chroma
from langchain.storage import InMemoryStore
from langchain.retrievers import ParentDocumentRetriever
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_core.documents import Document
from docx import Document as DocxDocument
import torch, os

def load_docs_from_directory(directory_path):
    """
    遍历指定文件夹中的所有文件，读取并转换为docs格式
    """
    docs = []
    for filename in os.listdir(directory_path):
        if filename.endswith(".docx"):
            file_path = os.path.join(directory_path, filename)
            docs.extend(docx_to_documents(file_path))
    return docs


# def docx_to_documents(docx_path: str):
#     docx = DocxDocument(docx_path)
#     documents = []
#     for para in docx.paragraphs:
#         paragraph = para.text.strip().replace(" ", "")
#         if paragraph:
#             doc = Document(page_content=paragraph)
#             documents.append(doc)
#     return documents



def docx_to_documents(docx_path: str, chunk_size: int = 3000):
    docx = DocxDocument(docx_path)
    documents = []
    buffer = ""
    filename = os.path.basename(docx_path)

    for para in docx.paragraphs:
        text = para.text.strip().replace(" ", "")
        if text:
            buffer += text
            while len(buffer) >= chunk_size:
                chunk = buffer[:chunk_size]
                documents.append(Document(page_content=chunk, metadata={"filename": filename}))
                buffer = buffer[chunk_size:]

    if buffer:
        documents.append(Document(page_content=buffer, metadata={"filename": filename}))

    return documents


def doc_initialization(docs):
    chunk_size = 30
    chunk_overlap = 10

    EMBEDDING_DEVICE = "cuda" if torch.cuda.is_available() else "mps" if torch.backends.mps.is_available() else "cpu"
    embeddings = HuggingFaceEmbeddings(
        model_name='bge-m3',
        model_kwargs={'device': EMBEDDING_DEVICE},
        encode_kwargs={'normalize_embeddings': True}
    )

    vectordb = Chroma(
        embedding_function=embeddings,
        collection_name="docs_collection",
        collection_metadata={"hnsw:space": "cosine"}
    )

    text_splitter = RecursiveCharacterTextSplitter(chunk_size=chunk_size, chunk_overlap=chunk_overlap)
    store = InMemoryStore()

    retriever = ParentDocumentRetriever(
        vectorstore=vectordb,
        docstore=store,
        child_splitter=text_splitter,
        search_kwargs={"k": 4},
        search_type="similarity",
    )

    # 拆分文档时保留父文档内容作为metadata
    split_docs = []
    for doc in docs:
        splits = text_splitter.split_text(doc.page_content)
        for split in splits:
            split_doc = Document(
                page_content=split,
                metadata={
                    "full_doc": doc.page_content,
                    "filename": doc.metadata.get("filename", "unknown")
                }
            )
            split_docs.append(split_doc)


    # 分批添加,否则Chroma可能会报错
    batch_size = 5000
    for i in range(0, len(split_docs), batch_size):
        batch = split_docs[i:i + batch_size]
        retriever.add_documents(batch)

    return retriever


def extract_context_snippet(child_doc: Document, context_chars: int = 200):
    full_text = child_doc.metadata.get("full_doc", "")
    filename = child_doc.metadata.get("filename", "unknown")
    if not full_text:
        return Document(
            page_content=child_doc.page_content,
            metadata={"filename": filename}
        )

    start_idx = full_text.find(child_doc.page_content)
    if start_idx == -1:
        return Document(
            page_content=child_doc.page_content,
            metadata={"filename": filename}
        )

    end_idx = start_idx + len(child_doc.page_content)
    context_start = max(0, start_idx - context_chars)
    context_end = min(len(full_text), end_idx + context_chars)

    snippet = full_text[context_start:context_end]
    return Document(
        page_content=snippet,
        metadata={"filename": filename}
    )


import time

if __name__ == '__main__':
    directory_path = r"切片"

    start_time = time.time()

    docs = load_docs_from_directory(directory_path)

    # print(docs)
    print("文档数量：", len(docs))

    retriever = doc_initialization(docs=docs)

    end_time = time.time()
    print("运行时间: {:.6f} 秒".format(end_time - start_time))
    
    while True:
        query = input("\n请输入查询语句（输入 e 退出）：")
        if query.lower() == "e":
            break

        results = retriever.get_relevant_documents(query)

        print("\n===== 检索结果（含上下300字上下文） =====\n")
        formatted_results = [extract_context_snippet(doc) for doc in results]

        print("formatted_results:", formatted_results)

        for i, doc in enumerate(formatted_results, 1):
            print(f"\n--- 结果 {i} | 文件: {doc.metadata.get('filename', 'unknown')} ---\n")
            print(doc.page_content)
