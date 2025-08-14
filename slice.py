from docx import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
import os

def extract_paragraphs_from_docx(file_path):
    """
    从指定路径的 docx 文件中提取段落内容，返回一个段落列表。
    """
    doc = Document(file_path)
    paragraphs = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:  # 跳过空段落
            paragraphs.append(text)

    return paragraphs


def load_docx_text(docx_path: str) -> str:
    """
    从指定路径加载docx文件内容，并拼接为一个字符串。
    """
    if not os.path.exists(docx_path):
        raise FileNotFoundError(f"文件未找到: {docx_path}")
    
    doc = Document(docx_path)
    paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
    return '\n'.join(paragraphs)

def clean_text(text: str) -> str:
    """
    删除文本中的换行符和多余空格。
    """
    return text.replace('\n', '').replace('\r', '').strip()

def split_text(text: str, chunk_size: int = 500, chunk_overlap: int = 50):
    """
    使用 LangChain 的 RecursiveCharacterTextSplitter 进行文本切分。
    """
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap
    )
    return text_splitter.split_text(text)


def slice_pdf(docx_path):
    
    raw_text = load_docx_text(docx_path)
    cleaned_text = clean_text(raw_text)

    # print("🧹 清洗后的文本：")
    # print(cleaned_text[:200] + "..." if len(cleaned_text) > 200 else cleaned_text)

    chunks = split_text(cleaned_text, chunk_size=300, chunk_overlap=50)

    # print(f"\n📚 总共分成 {len(chunks)} 段：")
    # for i, chunk in enumerate(chunks, 1):
    #     print(f"[第{i}段] {chunk}\n")

    return chunks


def main(docx_path: str):
    """
    主函数：根据文件名前缀判断处理方式。
    如果文件名以 'pdf' 开头，调用 slice_pdf；
    否则，调用 extract_paragraphs_from_docx。
    返回对应处理结果。
    """
    filename = os.path.basename(docx_path).lower()
    
    if filename.startswith("pdf"):
        ans = slice_pdf(docx_path)
        return ans
    else:
        ans = extract_paragraphs_from_docx(docx_path)
        return ans


# 示例用法
if __name__ == "__main__":
    docx_path = r"D:\projects\sailisi\切片\pdf_9_汽车零部件行业供应链质量管理与探讨——基于A公司在供应链质量管理的实践.docx"  # 替换为你的文件路径
    result = main(docx_path)

    print(result[:20])
